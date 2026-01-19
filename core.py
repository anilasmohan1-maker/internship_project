# core.py

import torch
from transformers import AutoTokenizer, AutoModelForCausalLM
from nltk.corpus import stopwords
from collections import Counter
from jinja2 import Environment, BaseLoader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from templates import MODERN_TEMPLATE

# -----------------------
# NLTK SETUP
# -----------------------
nltk.download("punkt", quiet=True)
nltk.download("stopwords", quiet=True)

# -----------------------
# MODEL LOADING
# -----------------------
def load_model():
    tokenizer = AutoTokenizer.from_pretrained("distilgpt2")
    model = AutoModelForCausalLM.from_pretrained("distilgpt2")
    model.eval()
    return model, tokenizer

# -----------------------
# ATS OPTIMIZER
# -----------------------
# core.py

import re
from collections import Counter

class ATSOptimizer:
    def __init__(self):
        self.stopwords = {
            "and", "or", "the", "a", "an", "to", "in", "of", "for",
            "with", "on", "at", "by", "from", "is", "are", "as"
        }

    def extract_keywords(self, job_desc, top_n=20):
        if not job_desc:
            return []

        # Regex-based tokenization (NO NLTK)
        tokens = re.findall(r"\b[a-zA-Z]{2,}\b", job_desc.lower())

        # Remove stopwords
        filtered_tokens = [t for t in tokens if t not in self.stopwords]

        # Count word frequency
        freq = Counter(filtered_tokens)

        # Return top keywords
        return [word for word, _ in freq.most_common(top_n)]


# -----------------------
# AI GENERATOR
# -----------------------
class ContentGenerator:
    def __init__(self, model, tokenizer):
        self.model = model
        self.tokenizer = tokenizer

    def generate_summary(self, prompt, max_new_tokens=180):
        inputs = self.tokenizer(prompt, return_tensors="pt")
        with torch.no_grad():
            output = self.model.generate(**inputs, max_new_tokens=max_new_tokens)
        return self.tokenizer.decode(output[0], skip_special_tokens=True)

# -----------------------
# DOCUMENT GENERATOR
# -----------------------
class DocumentGenerator:
    def __init__(self):
        self.env = Environment(loader=BaseLoader())

    def render_html(self, profile, summary):
        template = self.env.from_string(MODERN_TEMPLATE)
        return template.render(profile=profile, summary=summary)

    def generate_docx(self, profile, summary, filename="resume.docx"):
        doc = Document()

        title = doc.add_paragraph(profile["name"])
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(24)

        doc.add_heading("Professional Summary", level=2)
        doc.add_paragraph(summary)

        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(profile["skills"]))

        doc.save(filename)

        return filename
