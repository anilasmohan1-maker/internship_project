import streamlit as st
import json
import zipfile
from datetime import datetime
from pathlib import Path
from collections import Counter
import os
import re

import nltk
from nltk.corpus import stopwords

from huggingface_hub import InferenceClient
from jinja2 import Environment, BaseLoader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from samples import SAMPLE_PROFILES

# -------------------------------------------------
# PAGE CONFIG
# -------------------------------------------------
st.set_page_config(
    page_title="AI Resume & Portfolio Builder",
    layout="wide"
)

# -------------------------------------------------
# NLTK SETUP
# -------------------------------------------------
@st.cache_resource
def setup_nltk():
    nltk.download("punkt", quiet=True)
    nltk.download("stopwords", quiet=True)

setup_nltk()

# -------------------------------------------------
# LOAD HF INFERENCE CLIENT (REAL MODEL)
# -------------------------------------------------
@st.cache_resource
def load_hf_client():
    HF_API_KEY = st.secrets.get("HF_API_KEY") or os.getenv("HF_API_KEY")
    if not HF_API_KEY:
        raise ValueError("HF_API_KEY not found. Set it in Streamlit secrets or environment.")

    return InferenceClient(
        model="HuggingFaceH4/zephyr-7b-beta",
        token=HF_API_KEY
    )


hf_client = load_hf_client()

# -------------------------------------------------
# ATS OPTIMIZER
# -------------------------------------------------
class ATSOptimizer:
    def __init__(self):
        self.stop_words = set(stopwords.words("english"))

    def extract_keywords(self, text, top_n=20):
        tokens = re.findall(r"\b[a-zA-Z]{4,}\b", text.lower())
        filtered = [w for w in tokens if w not in self.stop_words]
        return [w for w, _ in Counter(filtered).most_common(top_n)]

# -------------------------------------------------
# AI CONTENT GENERATOR (MISTRAL 7B)
# -------------------------------------------------
class ContentGenerator:
    def __init__(self, client):
        self.client = client

    def _generate(self, prompt, max_tokens=400):
        response = self.client.text_generation(
            prompt,
            max_new_tokens=max_tokens,
            temperature=0.4,
            top_p=0.9,
            repetition_penalty=1.1
        )
        return response.strip()

    def generate_summary(self, profile):
        prompt = f"""
[INST]
You are a professional resume writer.

Write a realistic, ATS-optimized professional summary (3–4 lines).

Candidate: {profile['name']}
Target Role: {profile['targets']['title']}
Skills: {", ".join(profile['skills'])}

Rules:
- Mention years of experience (estimate realistically)
- Include tools, technologies, or domains
- Avoid buzzwords and generic phrases
- Sound like a real resume written by a human
[/INST]
"""
        return self._generate(prompt, 180)

    def generate_bullets(self, exp, keywords):
        prompt = f"""
[INST]
Write real resume bullet points.

Role: {exp['title']}
Company: {exp['company']}
Context: {exp['description']}

Rules:
- 3–4 bullets
- Start each bullet with •
- Use action verbs
- Mention tools, frameworks, metrics, datasets
- ATS-optimized but human

Keywords: {", ".join(keywords)}
[/INST]
"""
        return self._generate(prompt, 250)

    def generate_cover_letter(self, profile):
        prompt = f"""
[INST]
Write a concise professional cover letter (3 paragraphs).

Candidate: {profile['name']}
Email: {profile['email']}
Phone: {profile['phone']}
LinkedIn: {profile['linkedin']}

Target Role: {profile['targets']['title']}
Company: {profile['targets']['company']}

Tone:
- Professional
- Confident
- Clear
- No AI buzzwords
[/INST]
"""
        return self._generate(prompt, 450)

# -------------------------------------------------
# DOCUMENT GENERATOR
# -------------------------------------------------
class DocumentGenerator:
    def __init__(self):
        self.env = Environment(loader=BaseLoader())
        self.template_path = Path("resume.html")

    def render_html(self, profile, summary, exp_with_bullets):
        template = self.env.from_string(
            self.template_path.read_text(encoding="utf-8")
        )
        return template.render(
            profile=profile,
            summary=summary,
            exp_with_bullets=exp_with_bullets
        )

    def generate_docx(self, profile, summary, exp_with_bullets):
        doc = Document()

        header = doc.add_paragraph(profile["name"])
        header.runs[0].font.size = Pt(24)
        header.runs[0].bold = True
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact = doc.add_paragraph(
            f"{profile['email']} | {profile['phone']} | {profile['linkedin']}"
        )
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading("Professional Summary", level=2)
        doc.add_paragraph(summary)

        doc.add_heading("Experience", level=2)
        for exp, bullets in exp_with_bullets:
            p = doc.add_paragraph()
            p.add_run(f"{exp['title']} | {exp['company']}").bold = True
            doc.add_paragraph(bullets)

        doc.add_heading("Education", level=2)
        for edu in profile["education"]:
            p = doc.add_paragraph()
            p.add_run(edu["degree"]).bold = True
            p.add_run(f" | {edu['institution']} | GPA: {edu['gpa']}")

        doc.add_heading("Skills", level=2)
        doc.add_paragraph(", ".join(profile["skills"]))

        path = "resume.docx"
        doc.save(path)
        return path

    def create_zip(self, html, cover_letter, profile):
        zip_name = f"portfolio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
        with zipfile.ZipFile(zip_name, "w") as z:
            z.writestr("resume.html", html)
            z.writestr("cover_letter.txt", cover_letter)
            z.writestr("profile.json", json.dumps(profile, indent=2))
        return zip_name

# -------------------------------------------------
# OBJECTS
# -------------------------------------------------
ats = ATSOptimizer()
generator = ContentGenerator(hf_client)
docgen = DocumentGenerator()

# -------------------------------------------------
# SIDEBAR UI
# -------------------------------------------------
st.sidebar.title("⚙️ Candidate Details")

user_name = st.sidebar.text_input("Full Name", "Anila R")
user_email = st.sidebar.text_input("Email", "anila.r@email.com")
user_phone = st.sidebar.text_input("Phone", "+91 98765 43210")
user_linkedin = st.sidebar.text_input("LinkedIn", "https://linkedin.com/in/anilar")

st.sidebar.markdown("---")

profile_key = st.sidebar.selectbox(
    "Choose Sample Profile",
    list(SAMPLE_PROFILES.keys()),
    index=0
)

profile = SAMPLE_PROFILES[profile_key].copy()
profile["name"] = user_name
profile["email"] = user_email
profile["phone"] = user_phone
profile["linkedin"] = user_linkedin

# -------------------------------------------------
# MAIN UI
# -------------------------------------------------
st.title("🤖 AI Resume & Portfolio Builder")

if st.button("✨ Generate Resume & Portfolio"):
    with st.spinner("Generating professional resume content..."):
        keywords = ats.extract_keywords(profile["targets"]["job_description"])
        summary = generator.generate_summary(profile)

        bullets = [
            generator.generate_bullets(exp, keywords)
            for exp in profile["experience"]
        ]

        cover_letter = generator.generate_cover_letter(profile)
        exp_with_bullets = list(zip(profile["experience"], bullets))

        resume_html = docgen.render_html(profile, summary, exp_with_bullets)
        docx_path = docgen.generate_docx(profile, summary, exp_with_bullets)
        zip_path = docgen.create_zip(resume_html, cover_letter, profile)

    tabs = st.tabs(["📄 Resume Preview", "✉️ Cover Letter", "⬇️ Downloads"])

    with tabs[0]:
        st.components.v1.html(resume_html, height=900, scrolling=True)

    with tabs[1]:
        st.text_area("Cover Letter", cover_letter, height=450)

    with tabs[2]:
        st.download_button("⬇️ Resume (HTML)", resume_html, "resume.html")
        st.download_button("⬇️ Resume (DOCX)", open(docx_path, "rb"), "resume.docx")
        st.download_button("⬇️ Portfolio (ZIP)", open(zip_path, "rb"), zip_path)

